import json
import os
import re
from pathlib import Path
from typing import Any, Literal, TypedDict

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph
from dotenv import load_dotenv
from langchain_core.prompts import ChatPromptTemplate
from langchain_community.vectorstores import Chroma
from langchain_core.output_parsers import PydanticOutputParser
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_groq import ChatGroq
from langgraph.graph import END, START, StateGraph
from pydantic import BaseModel, Field

from .ingest import COMPANY_VAULT_COLLECTION, TENDER_COLLECTION

CORE_DIR = Path(__file__).resolve().parent
BACKEND_DIR = CORE_DIR.parent
ENV_FILE = BACKEND_DIR / ".env"
CHROMA_DIR = BACKEND_DIR / "chroma_db"
EMBEDDING_MODEL = "sentence-transformers/all-MiniLM-L6-v2"
GROQ_MODEL = "llama-3.3-70b-versatile"
OUTPUT_DIR = BACKEND_DIR / "output"
FINAL_DOCX = OUTPUT_DIR / "Final_Proposal.docx"

_NAVY = RGBColor(0, 32, 96)
_WARNING_RED = RGBColor(192, 0, 0)


def _set_corporate_styles(doc: Document) -> None:
    """Arial body; navy bold headings for corporate proposal styling."""
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)

    for level in (1, 2, 3):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "Arial"
        hs.font.bold = True
        hs.font.color.rgb = _NAVY


def _add_inline_bold_runs(paragraph: Paragraph, text: str) -> None:
    """Split on **...** pairs; odd segments are bold."""
    if not text:
        return
    chunks = text.split("**")
    for i, chunk in enumerate(chunks):
        if not chunk:
            continue
        run = paragraph.add_run(chunk)
        if i % 2 == 1:
            run.bold = True


def _add_cover_page(doc: Document, prepared_by: str | None = None) -> None:
    """Professional cover: title, subtitle, attribution, then page break."""
    bidder = (prepared_by or "").strip() or "[Bidder Name]"
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(12)
    tr = title_p.add_run("TECHNICAL PROPOSAL")
    tr.bold = True
    tr.font.name = "Arial"
    tr.font.size = Pt(36)
    tr.font.color.rgb = _NAVY

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(24)
    sr = sub_p.add_run("Response to PPRA Tender Requirements")
    sr.font.name = "Arial"
    sr.font.size = Pt(14)
    sr.font.color.rgb = _NAVY

    prep_p = doc.add_paragraph()
    prep_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prep_p.paragraph_format.space_before = Pt(18)
    pr = prep_p.add_run(f"Prepared by: {bidder}")
    pr.font.name = "Arial"
    pr.font.size = Pt(11)

    gen_p = doc.add_paragraph()
    gen_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    gen_p.paragraph_format.space_after = Pt(12)
    gr = gen_p.add_run("Generated via TenderForge PK")
    gr.font.name = "Arial"
    gr.font.size = Pt(11)
    gr.italic = True

    doc.add_page_break()


def _append_markdown_lines(doc: Document, draft: str) -> None:
    """Parse draft line-by-line: markdown headings, bullets, inline **bold**."""
    for raw_line in draft.split("\n"):
        stripped = raw_line.strip()
        if not stripped:
            doc.add_paragraph()
            continue

        bullet_match = re.match(r"^[\*\-]\s+(.+)$", stripped)
        if bullet_match:
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_bold_runs(p, bullet_match.group(1).strip())
            continue

        if stripped.startswith("###"):
            text = stripped[3:].strip()
            p = doc.add_paragraph(style="Heading 3")
            _add_inline_bold_runs(p, text)
            continue
        if stripped.startswith("##"):
            text = stripped[2:].strip()
            p = doc.add_paragraph(style="Heading 2")
            _add_inline_bold_runs(p, text)
            continue
        if stripped.startswith("#"):
            text = stripped[1:].strip()
            p = doc.add_paragraph(style="Heading 1")
            _add_inline_bold_runs(p, text)
            continue

        p = doc.add_paragraph()
        _add_inline_bold_runs(p, stripped)


def _extract_bidder_name(company_context: str) -> str | None:
    """Heuristic: derive legal name from retrieved company profile text for the cover page."""
    if not company_context or not company_context.strip():
        return None
    text = company_context.strip()
    patterns = [
        re.compile(
            r"(?i)(?:company|firm|organization|legal|registered)\s+name"
            r"\s*[:：\-\u2013]\s*([^\n]+?)(?:\n|$)"
        ),
        re.compile(r"(?i)bidder\s*[:：]\s*([^\n]+)"),
        re.compile(
            r"(?i)(?:proponent|applicant|contractor)\s*[:：]\s*([^\n]+)"
        ),
    ]
    for pat in patterns:
        m = pat.search(text)
        if m:
            name = m.group(1).strip()
            if 2 < len(name) < 200:
                return name
    for line in text.splitlines():
        s = line.strip()
        if 4 < len(s) < 100 and s.isupper():
            return s.title()
    return None


class TenderState(TypedDict, total=False):
    tender_context: str
    requirements_json: dict[str, Any]
    company_context: str
    is_valid_context: bool
    rejection_reason: str
    draft_proposal: str
    audit_feedback: str
    audit_status: str
    revision_count: int
    output_docx_path: str


def _state_str(state: TenderState, key: str, default: str = "") -> str:
    """TypedDict .get default is ignored when the key exists with value None."""
    v = state.get(key, default)
    return default if v is None else str(v)


def _state_dict(state: TenderState, key: str) -> dict[str, Any]:
    v = state.get(key)
    if v is None or not isinstance(v, dict):
        return {}
    return v


def _state_int(state: TenderState, key: str, default: int = 0) -> int:
    v = state.get(key, default)
    if v is None:
        return default
    try:
        return int(v)
    except (TypeError, ValueError):
        return default


class RequirementItem(BaseModel):
    requirement: str = Field(description="Name of requirement")
    evidence_needed: str = Field(description="Required documentary evidence")
    status: str = Field(
        description="One of: required, optional, unclear",
        pattern="^(required|optional|unclear)$",
    )
    notes: str = Field(description="Practical compliance guidance")


class RequirementsChecklist(BaseModel):
    checklist: list[RequirementItem] = Field(
        description="Checklist of eligibility and technical requirements"
    )
    critical_deadlines: list[str] = Field(
        default_factory=list,
        description="Submission and validity deadlines mentioned in tender",
    )
    missing_information: list[str] = Field(
        default_factory=list,
        description="Information that is not found in retrieved context",
    )


class AuditResult(BaseModel):
    status: Literal["PASS", "FAIL"] = Field(
        description="PASS if draft fully satisfies requirements_json; FAIL otherwise"
    )
    feedback: str = Field(
        description=(
            "If FAIL: list missing clauses and hallucinated or unsupported claims. "
            "If PASS: exactly 'All requirements met'"
        )
    )


class GatekeeperLLMResult(BaseModel):
    """LLM JSON uses is_valid + reason; mapped to TenderState keys in the node."""

    is_valid: bool = Field(
        description="True only if both documents appear authentic and relevant"
    )
    reason: str = Field(
        description="If is_valid is false, a concise explanation; otherwise brief confirmation"
    )


def _load_env() -> None:
    load_dotenv(dotenv_path=ENV_FILE)
    if not os.getenv("GROQ_API_KEY"):
        raise EnvironmentError(
            "Missing GROQ_API_KEY. Add it to your environment or .env file."
        )


def get_embeddings() -> HuggingFaceEmbeddings:
    return HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL)


def load_vector_store(
    persist_directory: Path = CHROMA_DIR,
    collection_name: str = TENDER_COLLECTION,
) -> Chroma:
    if not persist_directory.exists():
        raise FileNotFoundError(
            f"Chroma directory not found: {persist_directory}. Run ingest first."
        )

    return Chroma(
        collection_name=collection_name,
        persist_directory=str(persist_directory),
        embedding_function=get_embeddings(),
    )


def get_llm(temperature: float = 0.0) -> ChatGroq:
    _load_env()
    return ChatGroq(
        model_name=GROQ_MODEL,
        temperature=temperature,
    )


def _retrieve_context(query: str, k: int = 6) -> str:
    vector_store = load_vector_store()
    retriever = vector_store.as_retriever(search_kwargs={"k": k})
    docs = retriever.invoke(query)

    if not docs:
        return ""

    return "\n\n".join(doc.page_content for doc in docs)


def _build_company_rag_query(requirements: dict[str, Any]) -> str:
    lines: list[str] = []
    raw_list = requirements.get("checklist")
    if isinstance(raw_list, list):
        for item in raw_list:
            if isinstance(item, dict):
                req = str(item.get("requirement", "")).strip()
                if req:
                    lines.append(req)
    joined = "\n".join(f"- {ln}" for ln in lines)
    if not joined.strip():
        return (
            "Full company legal name, PEC license category, address, years in business, "
            "past performance on similar work, key personnel, certifications, and "
            "technical capacity for public-sector ICT or civil works."
        )
    return (
        f"Tender requirements to substantiate with company evidence:\n{joined}\n\n"
        "Retrieve: PEC registration, company legal name, prior contracts, project "
        "references, financial position, key staff, methodology, and QA where relevant."
    )


def _retrieve_company_context(requirements: dict[str, Any], k: int = 8) -> str:
    try:
        vector_store = load_vector_store(collection_name=COMPANY_VAULT_COLLECTION)
    except (FileNotFoundError, OSError):
        return ""
    retriever = vector_store.as_retriever(search_kwargs={"k": k})
    query = _build_company_rag_query(requirements)
    docs = retriever.invoke(query)
    if not docs:
        return ""
    return "\n\n".join(doc.page_content for doc in docs)


def _retrieve_gatekeeper_company_sample(k: int = 6) -> str:
    """Broad retrieval from company vault for gatekeeper (no checklist yet)."""
    return _retrieve_company_context({}, k=k)


def gatekeeper_node(state: TenderState) -> dict[str, Any]:
    """
    Reject irrelevant tender or company PDFs before the main pipeline.
    If API already ran sync validation, state contains is_valid_context True — skip LLM.
    """
    if state.get("is_valid_context") is True:
        return {}

    tender_ctx = _retrieve_context(
        query=(
            "Government procurement tender notice scope of work eligibility "
            "submission deadline PPRA bidding documents BOQ technical specifications"
        ),
        k=8,
    )
    company_ctx = _retrieve_gatekeeper_company_sample(k=8)

    if not tender_ctx.strip():
        return {
            "is_valid_context": False,
            "rejection_reason": (
                "The tender PDF produced no retrievable text — it may be blank, "
                "corrupted, or not a usable procurement document."
            ),
        }
    if not company_ctx.strip():
        return {
            "is_valid_context": False,
            "rejection_reason": (
                "The company profile PDF produced no retrievable text — it may be blank "
                "or not a usable company document."
            ),
        }

    parser = PydanticOutputParser(pydantic_object=GatekeeperLLMResult)
    prompt = ChatPromptTemplate.from_messages(
        [
            (
                "system",
                (
                    "You are a strict Document Authenticator. Review the provided Tender "
                    "Context and Company Context. If the Tender is not a genuine "
                    "government or corporate procurement/RFP document, OR if the Company "
                    "Context is clearly a student assignment, fictional story, or "
                    "irrelevant text, you MUST return is_valid: false and state exactly "
                    "why in the reason field. Otherwise, return true. Output strict JSON "
                    "only.\n{format_instructions}"
                ),
            ),
            (
                "human",
                (
                    "Tender Context (sample):\n{tender_ctx}\n\n"
                    "Company Context (sample):\n{company_ctx}\n\n"
                    "Respond with JSON only."
                ),
            ),
        ]
    )

    llm = get_llm(temperature=0.0)
    chain = prompt | llm | parser
    result = chain.invoke(
        {
            "tender_ctx": tender_ctx[:12000],
            "company_ctx": company_ctx[:12000],
            "format_instructions": parser.get_format_instructions(),
        }
    )

    return {
        "is_valid_context": result.is_valid,
        "rejection_reason": result.reason,
    }


def route_gatekeeper(state: TenderState) -> str:
    if state.get("is_valid_context") is True:
        return "extractor_node"
    return END


def extractor_node(state: TenderState) -> dict[str, Any]:
    query = (
        "Extract all eligibility and technical requirements from this tender, "
        "including PEC licenses, past experience, financial capacity, and "
        "mandatory compliance items."
    )
    context = _retrieve_context(query=query, k=8)

    if not context.strip():
        return {
            "tender_context": "",
            "requirements_json": {
                "checklist": [],
                "critical_deadlines": [],
                "missing_information": [
                    "No retrievable tender context found in local ChromaDB."
                ],
            },
        }

    parser = PydanticOutputParser(pydantic_object=RequirementsChecklist)
    prompt = ChatPromptTemplate.from_messages(
        [
            (
                "system",
                (
                    "You are a PPRA tender compliance analyst for Pakistan public "
                    "procurement. Extract only what is explicitly supported by the "
                    "provided context. Produce strict JSON only and do not add prose.\n"
                    "{format_instructions}"
                ),
            ),
            (
                "human",
                (
                    "Tender Context:\n{context}\n\n"
                    "Return a strict JSON checklist of eligibility and technical "
                    "requirements for bidder qualification."
                ),
            ),
        ]
    )

    llm = get_llm(temperature=0.0)
    chain = prompt | llm | parser
    structured = chain.invoke(
        {
            "context": context,
            "format_instructions": parser.get_format_instructions(),
        }
    )

    return {
        "tender_context": context,
        "requirements_json": structured.model_dump(),
    }


def generator_node(state: TenderState) -> dict[str, Any]:
    requirements = _state_dict(state, "requirements_json")
    audit_feedback = _state_str(state, "audit_feedback", "")
    revision_count = _state_int(state, "revision_count", 0)
    company_context = _retrieve_company_context(requirements, k=8)

    prompt = ChatPromptTemplate.from_messages(
        [
            (
                "system",
                (
                    "You are drafting a PPRA public procurement bid. Use the provided "
                    "Company Context to prove the bidder can meet the requirements. "
                    "Cite only facts supported by the Company Context and the requirements "
                    "list; do not fabricate. NEVER invent bracket placeholders such as "
                    "[Company Name] or [PEC Number]. If a specific requirement is not "
                    "supported by the Company Context, state clearly that the requirement "
                    "cannot be substantiated with the information on file, and do not "
                    "pretend the evidence exists. Use clear section headings, Pakistan "
                    "PPRA-style formality, and compliance language. If audit feedback "
                    "exists, revise the draft to close those gaps without inventing facts."
                ),
            ),
            (
                "human",
                (
                    "Requirements JSON:\n{requirements_json}\n\n"
                    "Company Context (RAG from company vault; only use what appears here for "
                    "corporate and capability claims):\n{company_context}\n\n"
                    "Audit Feedback (if any):\n{audit_feedback}\n\n"
                    "Current Revision Count: {revision_count}\n\n"
                    "Write a complete technical proposal that maps each key requirement to "
                    "verifiable company facts where available."
                ),
            ),
        ]
    )

    llm = get_llm(temperature=0.3)
    chain = prompt | llm
    company_block = company_context or (
        "No company profile was retrieved. Do not make up corporate facts; for each "
        "such gap, state that the requirement cannot be met or documented from the profile."
    )
    response = chain.invoke(
        {
            "requirements_json": json.dumps(requirements, indent=2),
            "company_context": company_block,
            "audit_feedback": audit_feedback or "No audit feedback yet.",
            "revision_count": revision_count,
        }
    )

    return {
        "draft_proposal": response.content,
        "company_context": company_context,
        "revision_count": revision_count + 1,
    }


def auditor_node(state: TenderState) -> dict[str, Any]:
    draft = _state_str(state, "draft_proposal", "").strip()
    requirements = _state_dict(state, "requirements_json")
    current_rc = _state_int(state, "revision_count", 0)

    parser = PydanticOutputParser(pydantic_object=AuditResult)
    prompt = ChatPromptTemplate.from_messages(
        [
            (
                "system",
                (
                    "You are a strict PPRA tender evaluator for Pakistan public "
                    "procurement. Compare the Technical Proposal draft strictly "
                    "against the requirements checklist JSON (the tender "
                    "obligations). "
                    "PASS only if every required checklist item is explicitly and "
                    "correctly addressed without fabrication. "
                    "FAIL if any required item is missing, vague, or if the draft "
                    "claims facts not supported by the checklist or typical "
                    "bidder-held documents. "
                    "Output strict JSON only — no markdown, no commentary.\n"
                    "{format_instructions}"
                ),
            ),
            (
                "human",
                (
                    "Requirements checklist (JSON):\n{requirements_json}\n\n"
                    "Technical Proposal draft:\n{draft}\n\n"
                    "Respond with JSON only."
                ),
            ),
        ]
    )

    llm = get_llm(temperature=0.0)
    chain = prompt | llm | parser
    result = chain.invoke(
        {
            "requirements_json": json.dumps(requirements, indent=2),
            "draft": draft or "(empty draft)",
            "format_instructions": parser.get_format_instructions(),
        }
    )

    return {
        "audit_feedback": result.feedback,
        "audit_status": result.status,
        "revision_count": current_rc + 1,
    }


def export_node(state: TenderState) -> dict[str, Any]:
    draft = _state_str(state, "draft_proposal", "")
    revision_count = _state_int(state, "revision_count", 0)
    audit_status = _state_str(state, "audit_status", "")
    company_context = _state_str(state, "company_context", "")
    bidder_name = _extract_bidder_name(company_context)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _set_corporate_styles(doc)
    _add_cover_page(doc, prepared_by=bidder_name)

    if audit_status == "FAIL" and revision_count >= 3:
        warning = doc.add_paragraph()
        warn_run = warning.add_run(
            "WARNING: MAXIMUM REVISIONS REACHED. MANUAL REVIEW REQUIRED."
        )
        warn_run.bold = True
        warn_run.font.color.rgb = _WARNING_RED
        warn_run.font.name = "Arial"
        doc.add_paragraph()

    _append_markdown_lines(doc, draft)

    out_path = _state_str(state, "output_docx_path", "") or None
    target = Path(out_path) if out_path else FINAL_DOCX
    target.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(target))

    return {}


def route_audit(state: TenderState) -> str:
    status = _state_str(state, "audit_status", "")
    revision_count = _state_int(state, "revision_count", 0)

    if status == "PASS":
        return "export_node"
    if status == "FAIL":
        if revision_count < 3:
            return "generator_node"
        return "export_node"
    return "export_node"


def build_graph():
    graph = StateGraph(TenderState)
    graph.add_node("gatekeeper_node", gatekeeper_node)
    graph.add_node("extractor_node", extractor_node)
    graph.add_node("generator_node", generator_node)
    graph.add_node("auditor_node", auditor_node)
    graph.add_node("export_node", export_node)
    graph.add_edge(START, "gatekeeper_node")
    graph.add_conditional_edges(
        "gatekeeper_node",
        route_gatekeeper,
        {
            "extractor_node": "extractor_node",
            END: END,
        },
    )
    graph.add_edge("extractor_node", "generator_node")
    graph.add_edge("generator_node", "auditor_node")
    graph.add_conditional_edges(
        "auditor_node",
        route_audit,
        {
            "export_node": "export_node",
            "generator_node": "generator_node",
        },
    )
    graph.add_edge("export_node", END)
    return graph.compile()


app = build_graph()


if __name__ == "__main__":
    initial: TenderState = {
        "tender_context": "",
        "requirements_json": {},
        "company_context": "",
        "draft_proposal": "",
        "audit_feedback": "",
        "audit_status": "",
        "revision_count": 0,
    }
    print("Running TenderForge PK graph...\n")
    for step in app.stream(initial):
        for node_name, node_output in step.items():
            print(f"--- Step: {node_name} ---")
            print(json.dumps(node_output, indent=2, default=str))
            print()
    print("Done.")

